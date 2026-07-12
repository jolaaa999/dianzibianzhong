"""Markdown → Word 文档转换脚本"""
import re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "USER_MANUAL.md"
DST = "USER_MANUAL.docx"

doc = Document()

# 页面边距
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = '微软雅黑'
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # 标题
    if line.startswith('# ') and not line.startswith('## '):
        p = doc.add_heading(line[2:], level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)

    # 表格
    elif line.startswith('|') and '|' in line:
        rows = [line]
        i += 1
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append(lines[i].strip())
            i += 1
        i -= 1
        # 跳过分隔行
        data = [r for r in rows if not re.match(r'^\|[:\-\s|]+\|$', r)]
        if len(data) < 2:
            i += 1; continue
        headers = [c.strip() for c in data[0].split('|')[1:-1]]
        ncols = len(headers)
        table = doc.add_table(rows=len(data), cols=ncols, style='Light Grid Accent 1')
        for ri, row in enumerate(data):
            cells = [c.strip() for c in row.split('|')[1:-1]]
            for ci, cell_text in enumerate(cells[:ncols]):
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                        run.font.name = '微软雅黑'
        doc.add_paragraph()

    # 代码块
    elif line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i].rstrip())
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(50, 50, 50)

    # 引用块 >
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.italic = True

    # 分割线
    elif line.startswith('---'):
        doc.add_paragraph('─' * 60)

    # 空行
    elif line.strip() == '':
        pass

    # 普通段落
    else:
        text = line
        p = doc.add_paragraph()
        # 处理 **粗体** 和 `代码`
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
            else:
                p.add_run(part)

    i += 1

doc.save(DST)
print(f'Done: {DST}')
